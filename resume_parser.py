# -*- coding: utf-8 -*-
"""简历解析模块：将上传的 PDF / DOCX / TXT 文件内容转换为纯文本。"""

from io import BytesIO

import docx
from pypdf import PdfReader

# 解析后文本的最短有效长度（字符数），用于识别扫描件或空文件
MIN_TEXT_LENGTH = 50


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """根据文件后缀解析简历内容，返回纯文本。

    Args:
        file_bytes: 上传文件的字节流。
        filename: 上传文件名，用于判断文件类型。

    Returns:
        简历的纯文本内容。

    Raises:
        ValueError: 文件类型不支持、解析失败或文本过短时抛出。
    """
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix == "pdf":
        text = _parse_pdf(file_bytes)
    elif suffix == "docx":
        text = _parse_docx(file_bytes)
    elif suffix == "txt":
        text = _parse_txt(file_bytes)
    else:
        raise ValueError("暂不支持的文件类型，请上传 PDF、DOCX 或 TXT 格式的简历。")

    text = text.strip()
    if len(text) < MIN_TEXT_LENGTH:
        raise ValueError(
            "简历可能是扫描件或内容过少，请改为文字版简历，"
            "或直接在“粘贴 JD”处粘贴简历文字。"
        )
    return text


def _parse_pdf(file_bytes: bytes) -> str:
    """使用 pypdf 的 PdfReader 解析 PDF 文件，逐页提取文本。"""
    try:
        reader = PdfReader(BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError(f"PDF 解析失败：{exc}，请确认文件未损坏且未加密。") from exc


def _parse_docx(file_bytes: bytes) -> str:
    """使用 python-docx 解析 DOCX 文件，包含段落与表格内容。"""
    try:
        document = docx.Document(BytesIO(file_bytes))
        lines = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines)
    except Exception as exc:
        raise ValueError(f"DOCX 解析失败：{exc}，请确认文件未损坏。") from exc


def _parse_txt(file_bytes: bytes) -> str:
    """解析 TXT 文件，优先使用 UTF-8，失败后回退到 GBK。"""
    for encoding in ("utf-8", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 文件编码无法识别，请使用 UTF-8 或 GBK 编码。")

